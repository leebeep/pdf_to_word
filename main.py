## 不支持读取pdf 中的图片并识别为文字 
## 支持读取pdf中的文字转为word中的文字

import os
from configparser import ConfigParser
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document
import datetime

## 读取pdf并返回字符串（整个内容是个字符串）
def read_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()
        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()
        content = return_str.getvalue()
        return_str.close()
        return content

## 过滤字符-空格
def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))  #字符映射转换表表
    return content.translate(mpa)   #字符串中要过滤的字符列表（即空格）

## 文字内容存到word
def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):                           #以\n 分行
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))     #每行都去除空格
    doc.save(file_path)

## pdf 转 word
def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)      #读取pdf并返回文字，这里content 是个字符串，如果需要精准，则涉及到
    save_text_to_word(content, word_file_path)  #文字内容存到word

## 主函数
def main():
    starttime = datetime.datetime.now()
    #读取目录
    config_parser = ConfigParser()
    config_parser.read('config.cfg')
    config = config_parser['default']

    tasks = []
    with ProcessPoolExecutor(max_workers=int(config['max_worker'])) as executor:
        for file in os.listdir(config['pdf_folder']):  #获取文件列表
            extension_name = os.path.splitext(file)[1] #截取每个文件的格式
            if extension_name != '.pdf':
                continue
            file_name = os.path.splitext(file)[0]                          #文件名称
            pdf_file = config['pdf_folder'] + '/' + file                   #带取出的带路径的pdf文件名称
            word_file = config['word_folder'] + '/' + file_name + '.docx'  #要放置的带路径的word文件名称
            print('正在处理: ', file)
            result = executor.submit(pdf_to_word, pdf_file, word_file)     #pdf 转 word
            tasks.append(result)
    while True:
        exit_flag = True
        for task in tasks:
            if not task.done():
                exit_flag = False
        if exit_flag:
            print('完成')
            #exit(0) #无错误退出
            break
    endtime = datetime.datetime.now()
    print((endtime - starttime).seconds)    


if __name__ == '__main__':
    main()
